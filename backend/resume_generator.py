import google.generativeai as genai
import os
from docx import Document
# from docx2pdf import convert # Optional for PDF conversion

def tailor_resume_with_ai(job_description, resume_text):
    if not resume_text:
        return "Error: Could not extract text from resume to send to AI."
    if not job_description:
        return "Error: Job description is empty."

    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        print("Error: GOOGLE_API_KEY not found in environment variables.")
        return "Error: AI API Key is not configured. Please contact support."

    try:
        genai.configure(api_key=api_key)
    except Exception as e:
        print(f"Error configuring Gemini API: {e}")
        return f"Error configuring AI service: {e}"

    # Model selection - check Google's documentation for latest recommended models
    # 'gemini-1.0-pro' is a good general text model. 'gemini-1.5-pro-latest' might be newer/better.
    model_name = 'gemini-1.0-pro' # or 'gemini-1.5-pro-latest' if available and preferred
    try:
        model = genai.GenerativeModel(model_name)
    except Exception as e:
        print(f"Error initializing Gemini model '{model_name}': {e}")
        return f"Error initializing AI model: {e}"

    prompt = f"""
    You are an expert career coach and professional resume writer.
    Your task is to meticulously tailor the provided resume text to align with the given Job Description (JD).
    Your goal is to achieve a very high degree of relevance and incorporate keywords from the JD naturally.

    **Instructions:**
    1.  **Analyze the Job Description:** Identify key skills, responsibilities, qualifications, and company values.
    2.  **Analyze the Original Resume:** Understand the candidate's experience, skills, and achievements.
    3.  **Rewrite Experience/Responsibilities:**
        *   Rephrase bullet points in the 'Experience', 'Projects', or similar sections to directly address the requirements in the JD.
        *   Use strong action verbs and incorporate specific keywords and terminology from the JD.
        *   Quantify achievements with metrics whenever possible. If exact numbers aren't in the original resume, you can make reasonable, professional inferences or use placeholder formats like "[achieved X% improvement]" or "[managed Y projects]".
    4.  **Skills Section:** Ensure the skills section highlights skills mentioned in the JD that are also present or implied in the original resume. You may suggest adding relevant skills if clearly applicable based on the resume and JD.
    5.  **Summary/Objective (if present):** Briefly tailor the summary or objective to reflect the target role in the JD.
    6.  **Maintain Structure:** Preserve the overall structure of the original resume (e.g., Contact Info, Summary, Experience, Education, Skills). Do NOT invent new sections unless absolutely necessary and a standard resume component.
    7.  **Professional Tone:** The output must be professional, concise, and impactful.
    8.  **Output Format:** Provide ONLY the full text of the new, tailored resume. Do not include any introductory phrases like "Here is the tailored resume:", or any disclaimers, or explanations of your changes. Just the resume content itself, ready to be copied into a document. Ensure clear separation between sections (e.g., using common resume section headers like "Experience", "Education", "Skills").

    **Job Description (JD):**
    ---
    {job_description}
    ---

    **Original Resume Text:**
    ---
    {resume_text}
    ---

    **Tailored Resume Text (Full Content):**
    """

    generation_config = genai.types.GenerationConfig(
        # candidate_count=1, # Default
        # stop_sequences=None,
        max_output_tokens=8192, # Gemini 1.0 Pro has a large token limit, adjust if using other models or expecting very long outputs
        temperature=0.4,      # Lower temperature for more deterministic, factual output
        # top_p=0.95,
        # top_k=40
    )

    # Configure safety settings to be less restrictive if needed, but be mindful of policy.
    # Default is often BLOCK_MEDIUM_AND_ABOVE for most categories.
    # Test with BLOCK_ONLY_HIGH if defaults are too aggressive for resume/JD content.
    safety_settings = [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    ]

    try:
        print(f"Sending request to Gemini model: {model_name}")
        response = model.generate_content(
            prompt,
            generation_config=generation_config,
            safety_settings=safety_settings
        )

        # Detailed check for response validity and blocked content
        if not response.candidates:
            block_reason_msg = "Unknown (no candidates)"
            if response.prompt_feedback and response.prompt_feedback.block_reason:
                block_reason_msg = response.prompt_feedback.block_reason.name
            error_message = f"Error: AI content generation failed. The request might have been blocked. Reason: {block_reason_msg}."
            if hasattr(response, 'prompt_feedback') and response.prompt_feedback:
                 error_message += f" Prompt Feedback: {response.prompt_feedback}"
            print(error_message)
            return error_message

        # Accessing text through parts
        if not response.candidates[0].content.parts:
            error_message = "Error: AI response is empty or not in the expected format (no parts)."
            print(error_message)
            return error_message
        
        tailored_text = "".join(part.text for part in response.candidates[0].content.parts if hasattr(part, 'text'))
        
        if not tailored_text.strip():
            finish_reason = response.candidates[0].finish_reason.name if response.candidates[0].finish_reason else "UNKNOWN"
            error_message = f"Error: AI generated an empty response. Finish Reason: {finish_reason}."
            # Log more details if available, e.g., safety ratings
            if response.candidates[0].safety_ratings:
                error_message += f" Safety Ratings: {response.candidates[0].safety_ratings}"
            print(error_message)
            return error_message

        print("Successfully received response from Gemini.")
        return tailored_text.strip()

    except Exception as e:
        # This catches errors during the API call itself (network, auth, etc.)
        # or issues with processing the response object.
        error_type = type(e).__name__
        print(f"Gemini API call error ({error_type}): {e}")
        # For some google API errors, e.message might be more informative
        error_details = str(e)
        if hasattr(e, 'message'):
            error_details = e.message
        
        # Specific check for common API issues
        if "API key not valid" in error_details or "PERMISSION_DENIED" in error_details:
            return "Error: AI API Key is invalid or lacks permissions. Please check your key and API settings."
        if "Quota" in error_details or "RESOURCE_EXHAUSTED" in error_details:
            return "Error: AI API quota exceeded. Please check your usage limits or try again later."
        
        return f"Error during AI resume tailoring: {error_details}"


def create_docx_from_text(text_content, output_path):
    doc = Document()
    # Add a title (optional, could be filename or a generic title)
    # doc.add_heading('Tailored Resume', level=0)

    # Basic paragraph-based generation.
    # This could be improved if the AI provides markdown or section indicators.
    current_heading_level = 1 # Start with H1 for major sections
    for line in text_content.split('\n'):
        stripped_line = line.strip()
        if not stripped_line: # Skip empty lines, but add one paragraph for multiple blank lines
            if hasattr(create_docx_from_text, "last_line_empty") and create_docx_from_text.last_line_empty:
                continue # Avoid multiple empty paragraphs
            doc.add_paragraph('')
            create_docx_from_text.last_line_empty = True
            continue
        
        create_docx_from_text.last_line_empty = False

        # Simple heuristic for detecting section headers (e.g., all caps, few words, ends with no punctuation)
        # This is very basic and can be improved.
        # Example: "EXPERIENCE", "SKILLS", "EDUCATION"
        if stripped_line.isupper() and len(stripped_line.split()) < 4 and not stripped_line.endswith(('.', ':', ';')):
            try:
                doc.add_heading(stripped_line.title(), level=min(current_heading_level, 3)) # Use title case for headings
            except Exception as e_heading:
                print(f"Could not add heading '{stripped_line}': {e_heading}. Adding as paragraph.")
                doc.add_paragraph(stripped_line)

        # Heuristic for bullet points (lines starting with *, -, o, •)
        elif stripped_line.startswith(('* ', '- ', 'o ', '• ')):
            doc.add_paragraph(stripped_line[2:], style='ListBullet') # Remove bullet character before adding
        else:
            doc.add_paragraph(stripped_line)
    
    try:
        doc.save(output_path)
        print(f"DOCX saved to {output_path}")
        delattr(create_docx_from_text, "last_line_empty") # Clean up state
    except Exception as e:
        print(f"Error saving DOCX file {output_path}: {e}")
        delattr(create_docx_from_text, "last_line_empty") # Clean up state
        raise # Re-raise the exception to be caught by the caller

# Optional: Convert DOCX to PDF (requires LibreOffice or MS Word, can be problematic on servers)
# def convert_docx_to_pdf(docx_path, output_dir):
#     try:
#         pdf_filename = os.path.basename(docx_path).replace(".docx", ".pdf")
#         pdf_path = os.path.join(output_dir, pdf_filename)
#         print(f"Attempting to convert {docx_path} to {pdf_path}")
#         convert(docx_path, pdf_path) # This is the call to docx2pdf
#         print(f"PDF successfully saved to {pdf_path}")
#         return pdf_path
#     except Exception as e:
#         print(f"Error converting DOCX to PDF using docx2pdf: {e}")
#         print("Ensure LibreOffice or MS Word is installed and in PATH if using docx2pdf.")
#         print("Consider alternative PDF generation methods for server environments.")
#         return None