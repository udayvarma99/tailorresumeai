from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import tempfile
from docx import Document as DocxDocument # Renamed to avoid conflict if Document is used elsewhere
from PyPDF2 import PdfReader
import google.generativeai as genai
from dotenv import load_dotenv
import traceback

# Load environment variables from .env file
load_dotenv()

app = Flask(__name__)

# Configure CORS
CORS(app) # Allows all origins by default, refine for production

# Configure Gemini API
try:
    gemini_api_key = os.getenv("GOOGLE_API_KEY")
    if not gemini_api_key:
        raise ValueError("GOOGLE_API_KEY not found in environment variables. Please set it in the .env file.")
    genai.configure(api_key=gemini_api_key)
except ValueError as ve:
    print(f"Configuration Error: {ve}")
    # You might want to exit or handle this more gracefully if the app can't run without the API key
except Exception as e:
    print(f"An unexpected error occurred during Gemini configuration: {e}")


# --- Text Extraction Functions ---
def extract_text_from_pdf(filepath):
    text = ""
    try:
        reader = PdfReader(filepath)
        if not reader.pages:
            print(f"Warning: No pages found in PDF {filepath}.")
            return ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n" # Add newline between pages
        return text.strip()
    except Exception as e:
        print(f"Error reading PDF {filepath}: {e}")
        traceback.print_exc()
        raise ValueError(f"Could not extract text from PDF: {e}")

def extract_text_from_docx(filepath):
    text = ""
    try:
        doc = DocxDocument(filepath)
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error reading DOCX {filepath}: {e}")
        traceback.print_exc()
        raise ValueError(f"Could not extract text from DOCX: {e}")

def get_resume_text_from_file(filepath):
    filename = os.path.basename(filepath)
    ext = filename.split('.')[-1].lower()
    if ext == 'pdf':
        return extract_text_from_pdf(filepath)
    elif ext == 'docx':
        return extract_text_from_docx(filepath)
    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX are supported.")

# --- AI Tailoring Function ---
def tailor_resume_with_gemini(job_description, resume_text):
    if not resume_text:
        return "Error: Resume text is empty. Cannot process."
    if not job_description:
        return "Error: Job description is empty. Cannot process."

    # Model selection
    model_name = 'gemini-1.0-pro' # Or 'gemini-1.5-pro-latest' if available
    try:
        model = genai.GenerativeModel(model_name)
    except Exception as e:
        print(f"Error initializing Gemini model '{model_name}': {e}")
        traceback.print_exc()
        return f"Error: AI model initialization failed. {e}"

    prompt = f"""
    You are an expert career coach and professional resume writer.
    Your task is to meticulously tailor the provided resume text to align with the given Job Description (JD).
    Your goal is to achieve a very high degree of relevance and incorporate keywords from the JD naturally.

    **Instructions:**
    1.  **Analyze the Job Description:** Identify key skills, responsibilities, qualifications, and company values.
    2.  **Analyze the Original Resume:** Understand the candidate's experience, skills, and achievements.
    3.  **Rewrite Experience/Responsibilities:**
        *   Rephrase bullet points in the 'Experience', 'Projects', or similar sections to directly address the requirements in the JD.
        *   Use strong action verbs and incorporate specific keywords and terminology from the JD.
        *   Quantify achievements with metrics whenever possible. If exact numbers aren't in the original resume, you can make reasonable, professional inferences or use placeholder formats like "[achieved X% improvement]" or "[managed Y projects]".
    4.  **Skills Section:** Ensure the skills section highlights skills mentioned in the JD that are also present or implied in the original resume. You may suggest adding relevant skills if clearly applicable based on the resume and JD.
    5.  **Summary/Objective (if present):** Briefly tailor the summary or objective to reflect the target role in the JD.
    6.  **Maintain Structure:** Preserve the overall structure of the original resume (e.g., Contact Info, Summary, Experience, Education, Skills). Do NOT invent new sections unless absolutely necessary and a standard resume component.
    7.  **Professional Tone:** The output must be professional, concise, and impactful.
    8.  **Output Format:** Provide ONLY the full text of the new, tailored resume, ready to be copied into a document. Do not include any introductory phrases like "Here is the tailored resume:", or any disclaimers, or explanations of your changes. Ensure clear separation between sections (e.g., using common resume section headers like "Experience", "Education", "Skills").

    **Job Description (JD):**
    ---
    {job_description}
    ---

    **Original Resume Text:**
    ---
    {resume_text}
    ---

    **Tailored Resume Text (Full Content):**
    """

    generation_config = genai.types.GenerationConfig(
        max_output_tokens=8192,
        temperature=0.4,
    )
    safety_settings = [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    ]

    try:
        print(f"Sending request to Gemini model: {model_name}...")
        response = model.generate_content(
            prompt,
            generation_config=generation_config,
            safety_settings=safety_settings
        )

        if not response.candidates:
            block_reason_msg = "Unknown (no candidates)"
            if response.prompt_feedback and response.prompt_feedback.block_reason:
                block_reason_msg = response.prompt_feedback.block_reason.name
            error_message = f"Error: AI content generation failed. The request might have been blocked. Reason: {block_reason_msg}."
            if hasattr(response, 'prompt_feedback') and response.prompt_feedback:
                 error_message += f" Prompt Feedback: {response.prompt_feedback}"
            print(error_message)
            return error_message

        if not response.candidates[0].content.parts:
            error_message = "Error: AI response is empty or not in the expected format (no parts)."
            print(error_message)
            return error_message
        
        tailored_text = "".join(part.text for part in response.candidates[0].content.parts if hasattr(part, 'text'))
        
        if not tailored_text.strip():
            finish_reason = "UNKNOWN"
            if response.candidates[0].finish_reason:
                 finish_reason = response.candidates[0].finish_reason.name
            error_message = f"Error: AI generated an empty response. Finish Reason: {finish_reason}."
            if response.candidates[0].safety_ratings:
                error_message += f" Safety Ratings: {response.candidates[0].safety_ratings}"
            print(error_message)
            return error_message

        print("Successfully received response from Gemini.")
        return tailored_text.strip()

    except Exception as e:
        error_type = type(e).__name__
        print(f"Gemini API call error ({error_type}): {e}")
        traceback.print_exc()
        error_details = str(e)
        if "API key not valid" in error_details or "PERMISSION_DENIED" in error_details:
            return "Error: AI API Key is invalid or lacks permissions. Please check your key and API settings."
        if "Quota" in error_details or "RESOURCE_EXHAUSTED" in error_details:
            return "Error: AI API quota exceeded. Please check your usage limits or try again later."
        return f"Error during AI resume tailoring: {e}"

# --- DOCX Creation Function ---
def create_docx_from_text_content(text_content, output_path):
    doc = DocxDocument()
    # Heuristic state for preventing multiple empty paragraphs
    create_docx_from_text_content.last_line_empty = False 

    for line in text_content.split('\n'):
        stripped_line = line.strip()
        
        if not stripped_line:
            if not create_docx_from_text_content.last_line_empty:
                doc.add_paragraph('') 
                create_docx_from_text_content.last_line_empty = True
            continue # Skip processing further for empty lines
        
        create_docx_from_text_content.last_line_empty = False

        # Basic heading detection (all caps, few words)
        if stripped_line.isupper() and 1 < len(stripped_line.split()) < 5 and not stripped_line.endswith(('.', ':', ';', ',')):
            try:
                doc.add_heading(stripped_line.title(), level=1)
            except Exception as e_heading:
                print(f"Could not add heading '{stripped_line}': {e_heading}. Adding as paragraph.")
                doc.add_paragraph(stripped_line)
        # Bullet point detection
        elif stripped_line.startswith(('* ', '- ', '• ')):
            # Remove the bullet marker and leading space, then add with Word's bullet style
            item_text = stripped_line[2:]
            doc.add_paragraph(item_text, style='ListBullet')
        elif stripped_line.startswith('o '): # another common bullet
            item_text = stripped_line[2:]
            doc.add_paragraph(item_text, style='ListBullet')
        else:
            doc.add_paragraph(stripped_line)
    
    try:
        doc.save(output_path)
        print(f"DOCX saved to {output_path}")
    except Exception as e:
        print(f"Error saving DOCX file {output_path}: {e}")
        traceback.print_exc()
        raise
    finally:
        # Clean up state attribute
        if hasattr(create_docx_from_text_content, "last_line_empty"):
            delattr(create_docx_from_text_content, "last_line_empty")


# --- Flask Route ---
@app.route('/tailor_resume', methods=['POST'])
def tailor_resume_route():
    if 'resume' not in request.files:
        return jsonify({"error": "No resume file part in the request."}), 400
    if 'job_description' not in request.form:
        return jsonify({"error": "No job_description part in the request."}), 400

    file = request.files['resume']
    jd_text = request.form.get('job_description', '').strip()

    if not jd_text:
        return jsonify({"error": "Job description cannot be empty."}), 400
    if file.filename == '':
        return jsonify({"error": "No resume file selected."}), 400

    # Validate file extension (optional but good practice)
    allowed_extensions = {'pdf', 'docx'}
    filename = secure_filename(file.filename)
    file_ext = filename.split('.')[-1].lower()
    if file_ext not in allowed_extensions:
        return jsonify({"error": "Invalid file type. Only PDF and DOCX are allowed."}), 400

    # Use a temporary directory that cleans up automatically
    with tempfile.TemporaryDirectory() as tmpdirname:
        try:
            original_filepath = os.path.join(tmpdirname, filename)
            file.save(original_filepath)
            print(f"File '{filename}' saved temporarily to '{original_filepath}'")

            # 1. Extract text from resume
            print("Extracting text from resume...")
            resume_text = get_resume_text_from_file(original_filepath)
            if not resume_text: # Check if extraction yielded any text
                return jsonify({"error": "Could not extract any text from the resume. It might be image-based or empty."}), 400
            print("Resume text extracted successfully.")

            # 2. Tailor resume with AI
            print("Sending to AI for tailoring...")
            tailored_text_or_error = tailor_resume_with_gemini(jd_text, resume_text)
            
            if tailored_text_or_error.startswith("Error:"):
                 print(f"AI Tailoring Error: {tailored_text_or_error}")
                 user_error_message = tailored_text_or_error # Be more specific with AI errors
                 if "API Key" in tailored_text_or_error or "quota" in tailored_text_or_error.lower():
                     user_error_message = "An issue occurred with the AI service. Please try again later."
                 return jsonify({"error": user_error_message}), 500
            print("AI tailoring successful.")

            # 3. Create a new DOCX file
            tailored_docx_filename = f"tailored_{filename.rsplit('.',1)[0]}.docx"
            tailored_docx_path = os.path.join(tmpdirname, tailored_docx_filename)
            
            print(f"Creating tailored DOCX at {tailored_docx_path}...")
            create_docx_from_text_content(tailored_text_or_error, tailored_docx_path)
            print("Tailored DOCX created.")

            return send_file(
                tailored_docx_path,
                as_attachment=True,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                download_name=tailored_docx_filename # Use the generated name
            )

        except ValueError as ve: # Catch errors from text extraction or unsupported format
            print(f"ValueError: {ve}")
            traceback.print_exc()
            return jsonify({"error": str(ve)}), 400
        except Exception as e:
            print(f"An unexpected error occurred in /tailor_resume: {e}")
            traceback.print_exc()
            return jsonify({"error": "An unexpected server error occurred. Please try again."}), 500
        # tmpdirname and its contents are automatically cleaned up here

if __name__ == '__main__':
    # Get FLASK_DEBUG from .env, default to False if not set or invalid
    flask_debug = os.getenv('FLASK_DEBUG', 'False').lower() in ('true', '1', 't')
    print(f"Starting Flask app with debug={flask_debug} on http://127.0.0.1:5000")
    app.run(host='127.0.0.1', port=5000, debug=flask_debug)